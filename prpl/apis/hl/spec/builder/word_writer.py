
from docx import Document
from docx.styles.style import _ParagraphStyle
from docx.styles.style import _TableStyle
import logging
import os
from collections import OrderedDict


class WordWriter:
    """Microsoft Word specification writer for prpl HL-API.

    It generates a new word file document for the prpl HL-API based on the specified template.

    Example:
        # Setup an API.
        from prpl.apis.hl.com import API as HLAPI

        # Create list of objects.
        api_objects = []

        # Create list of response codes.
        api_response_codes = []

        # Create list of version.
        api_versions = []

        # Create API.
        api = HLAPI(api_objects, api_response_codes, api_versions)

        # Import module.
        from prpl.apis.hl.spec.builder import WordWriter as HLAPISpecWriter

        # Load API.
        writer = HLAPIWordWriter(api, 'specs/prpl HL-API ({}).docx'.format(self.api.get_version()))

        # Generate report.
        writer.build()

    """

    def __init__(self, api, file, template='specs/templates/prpl.docx'):
        """Initializes the specification writer.

        Args:
            api (prpl.apis.hl.com.api): API to be parsed.
            file (str): Target filename for the specification.
            template (str): Word template to be used.

        """

        self.api = api
        self.file = file
        self.document = Document(template)

        # Init logger.
        self.logger = logging.getLogger('WordWriter')

        # Load template styles.
        self.logger.debug('Styles - Started looking up template styles.')
        styles = self.document.styles

        # Find styles.
        self.prpl_cover_version_number_style = list(
            filter(lambda x: type(x) is _ParagraphStyle and x.name == 'prplCoverVersionNumber', styles))[0]

        self.prpl_table_text_style = list(
            filter(lambda x: type(x) is _ParagraphStyle and x.name == 'prplTableText', styles))[0]

        self.paragraph_style = list(filter(lambda x: type(x) is _ParagraphStyle and x.name == 'Normal', styles))[0]
        self.table_style = list(filter(lambda x: type(x) is _TableStyle and x.name == 'Table Grid', styles))[0]

        self.logger.debug('Styles - Finished looking up template styles.')

        # Remove old file.
        self.logger.debug('File - Removing previous report "{}".'.format(self.file))
        try:
            os.remove(self.file)
        except FileNotFoundError:
            pass
        self.logger.debug('File - Finished removing previous report "{}".'.format(self.file))

    def _append_table(self, headers, entries):
        """Creates a new table with the specified entries.

        It includes a header row and populates the remaining rows with the contents of entries based on the attributes
        specified on the headers.

        Args:
            headers (OrderedDict): Table headers used to create the first row and fetch the contents of each entry.
                The keys depict the attributes of the entries, and the values correspond to the header names.
            entries (list<object>): List of entries to be appended to the table.

        """

        # Create new table.
        t = self.document.add_table(rows=len(entries) + 1, cols=len(headers))

        # Assign a style.
        t.style = self.table_style

        # Writer header.
        for idx, header in enumerate(headers.values()):
            row_cells = t.rows[0].cells
            row_cells[idx].text = header

        # Append entries.
        for idx_entry, entry in enumerate(entries):
            row_cells = t.rows[idx_entry + 1].cells

            for idx_header, header in enumerate(headers.keys()):
                row_cells[idx_header].text = str(getattr(entry, header))

    def _update_cover(self):
        """Updates cover with API version number."""

        for paragraph in self.document.paragraphs:
            # Lookup version variable.
            if '$(VERSION)' in paragraph.text:
                # Replace text with version number.
                paragraph.text = 'Version {}'.format(self.api.get_version())
                paragraph.style = self.prpl_cover_version_number_style
                self.logger.debug('Cover - Updated with version {}.'.format(paragraph.text))
                break

    def _append_change_log(self):
        """Updates the release notes table."""

        cl_table = self.document.tables[0]

        for v in self.api.versions:
            cl_table.add_row()
            cl_table.rows[-1].cells[0].text = v.number
            cl_table.rows[-1].cells[1].text = v.date
            cl_table.rows[-1].cells[2].text = v.get_changes()
            self.logger.debug('ChangeLog - Appended version {}.'.format(v.number))

    def _append_return_codes(self):
        """Adds return codes section."""

        # Add chapter header.
        self.document.add_heading('Return Codes', level=1)

        # Write a chapter description.
        p = self.document.add_paragraph(
            'The following list of return codes is applicable to all objects and procedures.',
            style=self.paragraph_style)

        # Append response codes table.
        self._append_table(
            OrderedDict([('code', 'Code'), ('name', 'Name'), ('sample', 'Sample'), ('description', 'Description')]),
            self.api.response_codes)

    def _append_procedures(self):
        """Adds procedures section."""

        # Add heading.
        self.document.add_heading('Procedures', level=1)

        # Iterate through each object.
        for idx, obj in enumerate(self.api.objects):
            # Include page break between objects unless first and last.
            if 0 < idx < len(self.api.objects):
                self.document.add_page_break()

            # Object Header.
            self.document.add_heading(obj.name, 2)

            self.logger.debug('Procedures - Appended object "{}".'.format(obj.name))

            # Iterate through each procedure.
            for idx_proc, procedure in enumerate(obj.procedures):
                # Include page break between procedures unless first and last.
                if 0 < idx_proc < len(obj.procedures):
                    self.document.add_page_break()

                # Procedure Header.
                self.document.add_heading(procedure.name, 3)

                # Description.
                self.document.add_paragraph(procedure.description, style=self.paragraph_style)

                # Usage.
                self.document.add_heading('Usage', 4)

                request_body = ''
                if procedure.sample_request != '-':
                    request_body = ' "{}RequestBody{}"'.format('{', '}')

                self.document.add_paragraph(
                    'ubus call {} {}{}'.format(obj.name, procedure.name, request_body), style=self.paragraph_style)

                # Input.
                self.document.add_heading('Input', 4)

                args = list(filter(lambda x: x.is_input is True, procedure.fields))
                if len(args) == 0:
                    self.document.add_paragraph('N/A.', style=self.paragraph_style)
                else:
                    self._append_table(OrderedDict([('name', 'Name'),
                                                    ('description', 'Description'),
                                                    ('type', 'Type'),
                                                    ('is_required', 'Required'),
                                                    ('notes', 'Notes')]),
                                       args)

                # Output.
                self.document.add_heading('Output', 4)

                fields = list(filter(lambda x: x.is_output is True, procedure.fields))
                if len(fields) == 0:
                    self.document.add_paragraph('N/A.', style=self.paragraph_style)
                else:
                    self._append_table(OrderedDict([('name', 'Name'),
                                                    ('description', 'Description'),
                                                    ('type', 'Type'),
                                                    ('notes', 'Notes')]),
                                       fields)

                self.logger.debug('Procedures - Appended procedure "{}".'.format(procedure.name))

    def _append_events(self):
        """Adds events section.

        Generates one table for each object and skips objects with no events.

        """

        # Add page break.
        self.document.add_page_break()

        # Include chapter header.
        self.document.add_heading('Events', level=1)

        # Filter out objects with no events.
        objects_with_events = list(filter(lambda x: len(x.events) > 0, self.api.objects))

        # Iterate through each object.
        for idx, obj in enumerate(objects_with_events):
            # Include page break between objects unless first and last.
            if 0 < idx < len(objects_with_events):
                self.document.add_page_break()

            # Create object header.
            self.document.add_heading('{}'.format(obj.name), level=2)

            # Append events table.
            self._append_table(OrderedDict([('code', 'Code'),
                                            ('name', 'Name'),
                                            ('description', 'Description'),
                                            ('sample', 'Sample')]),
                               obj.events)

            self.logger.debug('Events - Added events for object "{}" with {} entries.'.format(obj.name,
                                                                                              len(obj.events)))

    def build(self):
        """Generates a Word file specification for the HL-API."""

        # Cover.
        self.logger.debug('Cover - Started writing.')
        self._update_cover()
        self.logger.debug('Cover - Finished.\n')

        # Add Change-Log.
        self.logger.debug('ChangeLog - Started writing.')
        self._append_change_log()
        self.logger.debug('ChangeLog - Finished.\n')

        # Add Return Codes.
        self.logger.debug('Response Codes - Started writing.')
        self._append_return_codes()
        self.logger.debug('Response - Finished.\n')

        # Add Objects.
        self.logger.debug('Procedures - Started writing.')
        self._append_procedures()
        self.logger.debug('Procedures - Finished.\n')

        # Add Events.
        self.logger.debug('Events - Started writing.')
        self._append_events()
        self.logger.debug('Events - Finished.\n')

        # Save file.
        self.logger.debug('File - Saving.')
        self.document.save(self.file)
        self.logger.debug('File - Finished.\n')
