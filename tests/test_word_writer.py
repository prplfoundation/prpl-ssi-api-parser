
import unittest2
from docx import Document
import os

from prpl.apis.hl.spec.builder import WordWriter as HLAPISpecWriter
from prpl.apis.hl.com import API as HLAPI
from prpl.apis.hl.com import ResponseCode as HLAPIResponseCode
from prpl.apis.hl.com import Version as HLAPIVersion
from prpl.apis.hl.com import Object as HLAPIObject
from prpl.apis.hl.com import Procedure as HLAPIProcedure
from prpl.apis.hl.com import Event as HLAPIEvent


class TestWordWriter(unittest2.TestCase):
    """Tests the 'prpl.apis.hl.spec.builder.WordWriter' component."""

    def setUp(self):
        """Test environment setup."""

        # Setup a HL-API Version.
        api_version = HLAPIVersion('3.5', '2018-04-13')
        api_version.change_list.append((1, 'Added new "foo" object.'))

        # Setup a HL-API Response Code.
        api_response_code = HLAPIResponseCode(
            'OK',
            'A well-formed call was performed and successfully processed.',
            '{"Header":{"Name":"OK"},"Body":{"Id":0,"Name":"Guest"}}')

        # Setup a HL-API Object.
        api_object = HLAPIObject(1, 'User.Accounts.{AccountId}', 'User Account')

        # Setup a HL-API Procedure.
        api_procedure = HLAPIProcedure(
            'Set',
            'Modifies the account.',
            '{"Name":"Admin"}',
            '{"Header":{"Name":"OK"}}')

        api_object.procedures.append(api_procedure)

        # Setup a HL-API Event.
        api_event = HLAPIEvent(
            '0',
            'USER_ACCOUNTS_ADDED',
            'Raised when a new account is added.',
            '{"Header":{"Code":1,"Name":"USER_ACCOUNTS_ADDED"},"Body":{"AccountId":"User.Accounts.2"}}')

        api_object.events.append(api_event)

        # Create list of objects.
        api_objects = [api_object]

        # Create list of response codes.
        api_response_codes = [api_response_code]

        # Create list of version.
        api_versions = [api_version]

        # Setup an API.
        self.api = HLAPI(api_objects, api_response_codes, api_versions)

        # Setup the WordWriter.
        self.test_filename = 'tests/test-api.docx'
        self.spec_writer = HLAPISpecWriter(self.api, self.test_filename)
        self.spec_writer.build()

        self.document = Document(self.test_filename)

    def tearDown(self):
        """Test environment teardown."""

        os.remove(self.test_filename)

    def test__update_cover(self):
        """Tests the 'WordWriter._update_cover' method ability to update the cover with the API version."""

        paragraphs = list(map(lambda x: x.text, self.document.paragraphs))

        # Validate that '$(VERSION)' tag was replaced by the right version.
        self.assertNotIn('$(VERSION)', paragraphs)
        self.assertIn('Version 3.5', paragraphs)

    def test__append_change_log(self):
        """Tests the 'WordWriter._append_change_log' method ability to generate the change log table."""

        # Validate that the last change-log entry is correct.
        cl_table = self.document.tables[0]
        self.assertEqual(cl_table.rows[-1].cells[0].text, '3.5')
        self.assertEqual(cl_table.rows[-1].cells[1].text, '2018-04-13')
        self.assertEqual(cl_table.rows[-1].cells[2].text, '1. Added new "foo" object.')

    def test__append_return_codes(self):
        """Tests the 'WordWriter._append_return_codes' method ability to generate the return codes table."""

        # Validate that the last return code entry is correct.
        cl_table = self.document.tables[-3]
        self.assertEqual(cl_table.rows[-1].cells[0].text, 'OK')
        self.assertEqual(cl_table.rows[-1].cells[1].text, '{"Header":{"Name":"OK"},"Body":{"Id":0,"Name":"Guest"}}')
        self.assertEqual(cl_table.rows[-1].cells[2].text,
                         'A well-formed call was performed and successfully processed.')

    def test__append_procedures(self):
        """Tests the 'WordWriter._append_procedures' method ability to generate the list of procedures."""

        paragraphs = list(map(lambda x: x.text, self.document.paragraphs))

        self.assertIn('Modifies the account.', paragraphs)
        self.assertIn('ubus call User.Accounts.{AccountId} Set "{RequestBody}"', paragraphs)

    def test__append_events(self):
        """Tests the 'WordWriter._append_events' method ability to generate the events table."""

        # Validate that the last event entry is correct.
        cl_table = self.document.tables[-1]
        self.assertEqual(cl_table.rows[-1].cells[0].text, '0')
        self.assertEqual(cl_table.rows[-1].cells[1].text, 'USER_ACCOUNTS_ADDED')
        self.assertEqual(cl_table.rows[-1].cells[2].text, 'Raised when a new account is added.')
        self.assertEqual(cl_table.rows[-1].cells[3].text,
                         '{"Header":{"Code":1,"Name":"USER_ACCOUNTS_ADDED"},"Body":{"AccountId":"User.Accounts.2"}}')

    def test_build(self):
        """Tests the 'WordWriter.build' method ability to generate word document."""

        os.path.isfile(self.test_filename)


if __name__ == '__main__':
    unittest2.main()
